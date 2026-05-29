#!/usr/bin/env python3
"""
pdf_to_docx_worker.py — PDF inspect + 文字型 PDF 轉 DOCX + Vision rebuild pipeline。

Phase 1a:inspect(metadata + 加密/掃描)
Phase 1b:convert(pdf2docx 主路線)
Phase 2a:render-pages(每頁 PNG + 嵌入圖抽出,給 Node 端餵 Gemini Vision)
Phase 2c:build-docx(讀 vision JSON → python-docx 組原生 table + bg fill + merge)

CLI(全程 stdout 只吐一行 JSON,stderr 走 log):
  python pdf_to_docx_worker.py inspect      --in <pdf> [--password <pwd>]
  python pdf_to_docx_worker.py convert      --in <pdf> --out <docx> [--password <pwd>]
  python pdf_to_docx_worker.py render-pages --in <pdf> --out-dir <dir> [--password <pwd>] [--dpi 200]
  python pdf_to_docx_worker.py build-docx   --in-json <vision_json> --out <docx>

JSON 回傳一律含 "ok": bool
  ok=False:
    error_code: PASSWORD_REQUIRED | PASSWORD_WRONG | OPEN_FAILED | CONVERT_FAILED | INVALID_ARGS | BUILD_FAILED
    error: human-readable
  ok=True (inspect):
    pages, encrypted, scanned_ratio, text_pages, image_pages, complexity_score, file_size_bytes
  ok=True (convert):
    out_path, pages_converted, elapsed_ms
  ok=True (render-pages):
    out_dir, dpi, pages: [{ page_index, png_path, width, height,
                            embedded_images: [{img_path, bbox}],
                            text_dict_path }]
  ok=True (build-docx):
    out_path, blocks_built, elapsed_ms
"""
import argparse
import json
import os
import sys
import time
import traceback

# Windows 預設 stdout = cp950,中文 JSON 會炸 Node utf8 decoder。
# Python 3.7+ 支援 reconfigure(Docker Linux 上本來就 utf-8,無害)。
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

# 掃描頁判定閾值:可選取文字 < N 字 且 有至少 1 張圖
SCANNED_TEXT_THRESHOLD = 50
SCANNED_RATIO_THRESHOLD = 0.8  # 整份 ≥ 80% 頁是掃描頁 → 視為掃描型 PDF


def log(msg):
    """stderr 給 Node 那邊收 log,不污染 stdout JSON。"""
    print(f"[pdf_worker] {msg}", file=sys.stderr, flush=True)


def emit(payload):
    """stdout 只能有一行 JSON。"""
    print(json.dumps(payload, ensure_ascii=False), flush=True)


def open_pdf(path, password=None):
    """
    開檔 + 處理加密。
    Returns (doc, error_payload):
      成功 → (doc, None)
      加密缺密碼 / 密碼錯 → (None, {ok:False, error_code:...})
    """
    import fitz  # PyMuPDF

    try:
        doc = fitz.open(path)
    except Exception as e:
        return None, {
            "ok": False,
            "error_code": "OPEN_FAILED",
            "error": f"無法開啟 PDF:{e}",
        }

    if doc.needs_pass:
        if not password:
            doc.close()
            return None, {
                "ok": False,
                "error_code": "PASSWORD_REQUIRED",
                "error": "此 PDF 已加密,需要密碼",
            }
        # authenticate 回傳:0=失敗,>0=成功(user/owner 不同等級)
        if doc.authenticate(password) == 0:
            doc.close()
            return None, {
                "ok": False,
                "error_code": "PASSWORD_WRONG",
                "error": "PDF 密碼錯誤",
            }
    return doc, None


def cmd_inspect(args):
    doc, err = open_pdf(args.input, args.password)
    if err:
        emit(err)
        return

    try:
        total = doc.page_count
        text_pages = 0
        image_pages = 0
        scanned_pages = 0
        # 複雜度啟發式:drawings(rect path/line)多 + image 多 → 複雜表格 / 圖文混排
        # 給 LLM / skill 判斷是否該走 Vision rebuild mode 而非 pdf2docx editable
        drawings_total = 0
        images_total = 0

        for i in range(total):
            page = doc.load_page(i)
            text_len = len(page.get_text("text").strip())
            img_count = len(page.get_images(full=False))
            try:
                draw_count = len(page.get_drawings())
            except Exception:
                draw_count = 0
            if text_len >= SCANNED_TEXT_THRESHOLD:
                text_pages += 1
            if img_count > 0:
                image_pages += 1
            if text_len < SCANNED_TEXT_THRESHOLD and img_count >= 1:
                scanned_pages += 1
            drawings_total += draw_count
            images_total += img_count

        scanned_ratio = (scanned_pages / total) if total > 0 else 0.0
        # 簡化公式:每頁平均 drawings + 每頁平均 images*5
        # > 30:複雜(密集 table) → 建議 vision mode
        # > 80:非常複雜(每頁百來個 drawings) → 強烈建議 vision mode
        avg_drawings = drawings_total / total if total > 0 else 0
        avg_images = images_total / total if total > 0 else 0
        complexity_score = min(100, int(avg_drawings / 5 + avg_images * 5))

        try:
            file_size = os.path.getsize(args.input)
        except OSError:
            file_size = None

        emit({
            "ok": True,
            "pages": total,
            "encrypted": bool(doc.needs_pass),  # 已用密碼解開仍視為原本加密
            "scanned_ratio": round(scanned_ratio, 3),
            "is_scanned_pdf": scanned_ratio >= SCANNED_RATIO_THRESHOLD,
            "text_pages": text_pages,
            "image_pages": image_pages,
            "scanned_pages": scanned_pages,
            "drawings_total": drawings_total,
            "images_total": images_total,
            "complexity_score": complexity_score,
            "recommended_mode": "vision" if complexity_score >= 30 else "editable",
            "file_size_bytes": file_size,
        })
    finally:
        doc.close()


def cmd_render_pages(args):
    """
    Render PDF 每頁成 PNG + 抽出嵌入圖,給 Node 端餵 Gemini Vision。

    輸出 layout:
      <out_dir>/page_001.png         — 每頁 render(zoom=dpi/72)
      <out_dir>/page_001_text.json   — PyMuPDF 結構化文字 dict(給 vision 當 ground-truth 用)
      <out_dir>/img_001_a.png        — 該頁嵌入圖(若有,Phase 2c build-docx 會嵌回 docx)
    """
    import fitz

    out_dir = args.out_dir
    dpi = args.dpi or 200
    zoom = dpi / 72.0

    doc, err = open_pdf(args.input, args.password)
    if err:
        emit(err)
        return

    try:
        os.makedirs(out_dir, exist_ok=True)
        pages_out = []
        total = doc.page_count
        mat = fitz.Matrix(zoom, zoom)

        for i in range(total):
            page = doc.load_page(i)
            page_no = i + 1
            png_path = os.path.join(out_dir, f"page_{page_no:03d}.png")
            text_json_path = os.path.join(out_dir, f"page_{page_no:03d}_text.json")

            # Render PNG(alpha=False 省空間,table 文字辨識不需要透明)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            pix.save(png_path)
            width, height = pix.width, pix.height
            del pix  # 釋放 RAM

            # 結構化文字 dict — 給 vision 當 ground-truth 用,避免 OCR 中文錯字
            try:
                tdict = page.get_text("dict")
                # 只留必要欄位,避免 JSON 太大
                simplified = {
                    "page": page_no,
                    "width": tdict.get("width"),
                    "height": tdict.get("height"),
                    "blocks": [
                        {
                            "bbox": b.get("bbox"),
                            "lines": [
                                {
                                    "bbox": l.get("bbox"),
                                    "spans": [
                                        {"text": s.get("text"), "bbox": s.get("bbox")}
                                        for s in (l.get("spans") or [])
                                    ],
                                }
                                for l in (b.get("lines") or [])
                            ],
                        }
                        for b in (tdict.get("blocks") or [])
                        if b.get("type") == 0  # 0=text block, 1=image
                    ],
                }
                with open(text_json_path, "w", encoding="utf-8") as f:
                    json.dump(simplified, f, ensure_ascii=False)
            except Exception as e:
                log(f"page {page_no} text dict extract failed: {e}")
                text_json_path = None

            # 抽嵌入圖 — 簡化版只記 xref,Phase 2c build-docx 內再抽,避免 RAM 暴漲
            embedded = []
            try:
                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    embedded.append({"xref": xref})
            except Exception as e:
                log(f"page {page_no} get_images failed: {e}")

            pages_out.append({
                "page_index": i,
                "page_no": page_no,
                "png_path": png_path,
                "text_json_path": text_json_path,
                "width": width,
                "height": height,
                "embedded_image_count": len(embedded),
            })

        emit({
            "ok": True,
            "out_dir": os.path.abspath(out_dir),
            "dpi": dpi,
            "total_pages": total,
            "pages": pages_out,
        })
    finally:
        doc.close()


def cmd_convert(args):
    if not args.output:
        emit({"ok": False, "error_code": "INVALID_ARGS", "error": "convert 需要 --out"})
        return

    # 先 inspect 一次處理加密(open_pdf 已包好錯誤回傳)
    doc, err = open_pdf(args.input, args.password)
    if err:
        emit(err)
        return
    pages = doc.page_count
    doc.close()

    # 走 pdf2docx 主路線
    try:
        from pdf2docx import Converter
    except Exception as e:
        emit({"ok": False, "error_code": "CONVERT_FAILED", "error": f"pdf2docx import 失敗:{e}"})
        return

    t0 = time.time()
    try:
        # pdf2docx 自己會處理密碼(它內部也用 PyMuPDF);若 None 就傳 None
        cv = Converter(args.input, password=args.password) if args.password else Converter(args.input)
        try:
            cv.convert(args.output, start=0, end=None)
        finally:
            cv.close()
    except Exception as e:
        log(f"convert exception: {traceback.format_exc()}")
        emit({
            "ok": False,
            "error_code": "CONVERT_FAILED",
            "error": f"pdf2docx 轉換失敗:{e}",
        })
        return

    elapsed_ms = int((time.time() - t0) * 1000)
    emit({
        "ok": True,
        "out_path": os.path.abspath(args.output),
        "pages_converted": pages,
        "elapsed_ms": elapsed_ms,
    })


def _set_cell_bg(cell, hex_color):
    """python-docx 沒有直接 API 設 cell shading,走 XML。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    if not hex_color:
        return
    color = hex_color.lstrip("#")
    if len(color) != 6:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def _apply_run_format(run, *, bold=False, italic=False, color=None, size=None):
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        from docx.shared import RGBColor
        c = color.lstrip("#")
        if len(c) == 6:
            try:
                run.font.color.rgb = RGBColor.from_string(c)
            except Exception:
                pass
    if size:
        from docx.shared import Pt
        try:
            run.font.size = Pt(float(size))
        except Exception:
            pass


def _build_table(doc, table_block):
    """
    Vision JSON convention(dense row format):
      {
        "type": "table",
        "rows": [
          [ {text, bg, color, bold, italic, size, colspan, rowspan, align}, null, null, {...} ],
          [ {...}, {...}, {...}, {...} ]
        ]
      }
    被 merge 進的 cell 位置用 null 佔位。
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt

    rows = table_block.get("rows") or []
    if not rows:
        return
    n_rows = len(rows)
    n_cols = max((len(r) for r in rows), default=0)
    if n_cols == 0:
        return

    table = doc.add_table(rows=n_rows, cols=n_cols)
    # table 樣式預設無 border 也常見;若 vision 標 "with_border" 就套
    try:
        table.style = table_block.get("style") or "Table Grid"
    except KeyError:
        pass

    # 1) 寫文字 + bg + 字型(先寫,再 merge,merge 後 cell 索引會跑掉)
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            spec = row[c_idx] if c_idx < len(row) else None
            if spec is None:
                continue
            cell = table.cell(r_idx, c_idx)
            text = str(spec.get("text") or "")
            # 清空既有 paragraph(add_table 預設給每 cell 一個空 paragraph)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            _apply_run_format(
                run,
                bold=bool(spec.get("bold")),
                italic=bool(spec.get("italic")),
                color=spec.get("color"),
                size=spec.get("size"),
            )
            align = (spec.get("align") or "").lower()
            if align == "center":
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if spec.get("bg"):
                _set_cell_bg(cell, spec["bg"])

    # 2) Merge(從大到小避免 cell 索引衝突 — 但實際上 python-docx merge 後仍能用原 row/col 索引)
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            spec = row[c_idx] if c_idx < len(row) else None
            if spec is None:
                continue
            cs = int(spec.get("colspan") or 1)
            rs = int(spec.get("rowspan") or 1)
            if cs <= 1 and rs <= 1:
                continue
            try:
                end_r = min(n_rows - 1, r_idx + rs - 1)
                end_c = min(n_cols - 1, c_idx + cs - 1)
                table.cell(r_idx, c_idx).merge(table.cell(end_r, end_c))
            except Exception as e:
                log(f"merge ({r_idx},{c_idx}) span ({rs},{cs}) failed: {e}")


def cmd_build_docx(args):
    """
    讀 vision JSON → 用 python-docx 組原生 DOCX。

    Input JSON shape:
      {
        "pages": [
          {
            "page_no": 1,
            "blocks": [
              {"type":"heading","level":1,"text":"...","color":"#1a73e8"},
              {"type":"paragraph","text":"...","bold":false},
              {"type":"table","rows":[[{...},null], [{...},{...}]]},
              {"type":"image","path":"/abs/path.png","width_inches":5.5},
              {"type":"page_break"}
            ]
          },
          ...
        ]
      }
    """
    if not args.in_json:
        emit({"ok": False, "error_code": "INVALID_ARGS", "error": "build-docx 需要 --in-json"})
        return
    if not args.output:
        emit({"ok": False, "error_code": "INVALID_ARGS", "error": "build-docx 需要 --out"})
        return

    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as e:
        emit({"ok": False, "error_code": "BUILD_FAILED", "error": f"python-docx import 失敗:{e}"})
        return

    try:
        with open(args.in_json, "r", encoding="utf-8") as f:
            payload = json.load(f)
    except Exception as e:
        emit({"ok": False, "error_code": "INVALID_ARGS", "error": f"無法讀 vision JSON:{e}"})
        return

    pages = payload.get("pages") or []
    if not pages:
        emit({"ok": False, "error_code": "BUILD_FAILED", "error": "vision JSON 沒有 pages"})
        return

    t0 = time.time()
    blocks_built = 0
    try:
        doc = Document()
        for page_idx, page in enumerate(pages):
            blocks = page.get("blocks") or []
            for blk in blocks:
                btype = blk.get("type")
                try:
                    if btype == "heading":
                        level = int(blk.get("level") or 1)
                        h = doc.add_heading(str(blk.get("text") or ""), level=min(max(level, 1), 6))
                        if blk.get("color"):
                            for r in h.runs:
                                _apply_run_format(r, color=blk["color"])
                    elif btype == "paragraph":
                        p = doc.add_paragraph()
                        run = p.add_run(str(blk.get("text") or ""))
                        _apply_run_format(
                            run,
                            bold=bool(blk.get("bold")),
                            italic=bool(blk.get("italic")),
                            color=blk.get("color"),
                            size=blk.get("size"),
                        )
                    elif btype == "table":
                        _build_table(doc, blk)
                    elif btype == "image":
                        img_path = blk.get("path")
                        if img_path and os.path.isfile(img_path):
                            try:
                                w = blk.get("width_inches")
                                if w:
                                    doc.add_picture(img_path, width=Inches(float(w)))
                                else:
                                    doc.add_picture(img_path)
                            except Exception as e:
                                log(f"add_picture {img_path} failed: {e}")
                    elif btype == "page_break":
                        doc.add_page_break()
                    elif btype == "spacer":
                        doc.add_paragraph()
                    blocks_built += 1
                except Exception as e:
                    log(f"block {btype} failed: {e}\n{traceback.format_exc()}")
            # 每頁之間加 page break(最後一頁不加)
            if page_idx < len(pages) - 1:
                doc.add_page_break()

        doc.save(args.output)
    except Exception as e:
        log(f"build_docx exception: {traceback.format_exc()}")
        emit({"ok": False, "error_code": "BUILD_FAILED", "error": f"組裝 DOCX 失敗:{e}"})
        return

    emit({
        "ok": True,
        "out_path": os.path.abspath(args.output),
        "blocks_built": blocks_built,
        "elapsed_ms": int((time.time() - t0) * 1000),
    })


def main():
    parser = argparse.ArgumentParser(description="PDF → DOCX worker")
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_ins = sub.add_parser("inspect", help="檢查 PDF metadata / 加密 / 掃描")
    p_ins.add_argument("--in", dest="input", required=True)
    p_ins.add_argument("--password", default=None)

    p_cv = sub.add_parser("convert", help="文字型 PDF → DOCX (pdf2docx)")
    p_cv.add_argument("--in", dest="input", required=True)
    p_cv.add_argument("--out", dest="output", required=True)
    p_cv.add_argument("--password", default=None)

    p_rp = sub.add_parser("render-pages", help="每頁 render PNG + 抽嵌入圖,給 Node 餵 Gemini Vision")
    p_rp.add_argument("--in", dest="input", required=True)
    p_rp.add_argument("--out-dir", dest="out_dir", required=True)
    p_rp.add_argument("--password", default=None)
    p_rp.add_argument("--dpi", type=int, default=200)

    p_bd = sub.add_parser("build-docx", help="讀 vision JSON → python-docx 組原生 DOCX")
    p_bd.add_argument("--in-json", dest="in_json", required=True)
    p_bd.add_argument("--out", dest="output", required=True)

    args = parser.parse_args()

    # build-docx 不需要 PDF input;其他 sub-cmd 都需要 --in
    if args.cmd != "build-docx":
        if not os.path.isfile(args.input):
            emit({
                "ok": False,
                "error_code": "OPEN_FAILED",
                "error": f"檔案不存在:{args.input}",
            })
            return

    try:
        if args.cmd == "inspect":
            cmd_inspect(args)
        elif args.cmd == "convert":
            cmd_convert(args)
        elif args.cmd == "render-pages":
            cmd_render_pages(args)
        elif args.cmd == "build-docx":
            cmd_build_docx(args)
    except Exception as e:
        log(f"unhandled exception: {traceback.format_exc()}")
        emit({
            "ok": False,
            "error_code": "OPEN_FAILED",
            "error": f"worker uncaught: {e}",
        })


if __name__ == "__main__":
    main()
